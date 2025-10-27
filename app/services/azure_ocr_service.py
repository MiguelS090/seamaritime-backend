import logging
import time
from typing import Dict, Any, List, Optional
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from app.core.config import settings

logger = logging.getLogger(__name__)

class AzureOCRService:
    """Serviço de OCR usando Azure Document Intelligence"""
    
    def __init__(self):
        self.client = DocumentAnalysisClient(
            endpoint=settings.AZURE_FORM_RECOGNIZER_ENDPOINT,
            credential=AzureKeyCredential(settings.AZURE_FORM_RECOGNIZER_API_KEY)
        )
    
    def process_q88_document(self, file_path: str) -> Dict[str, Any]:
        """
        Processa documento Q88 usando Azure Document Intelligence.
        Retorna apenas tokens puros (linhas, parágrafos, tabelas) - sem extração de campos.
        """
        try:
            # Verificar se é .docx e extrair texto diretamente
            if file_path.lower().endswith('.docx'):
                return self._process_docx(file_path)
            
            with open(file_path, 'rb') as f:
                poller = self.client.begin_analyze_document(
                    "prebuilt-read", 
                    document=f
                )
                analyze_result = poller.result()
            
            # Extrair texto completo
            full_text = analyze_result.content if analyze_result.content else ""
            
            # Organizar linhas com informações de posição
            organized_lines = []
            for page in analyze_result.pages:
                for line in page.lines:
                    line_text = line.content.strip()
                    
                    organized_lines.append({
                        'text': line_text,
                        'confidence': 0.9,  # Valor padrão para linhas
                        'page_number': page.page_number,
                        'bounding_box': [point for point in line.polygon] if line.polygon else [],
                        'spans': line.spans
                    })
            
            # Extrair parágrafos
            paragraphs = []
            for paragraph in analyze_result.paragraphs:
                if paragraph.content.strip():
                    paragraph_data = {
                        'content': paragraph.content.strip(),
                        'confidence': 0.9,  # Valor padrão para parágrafos
                        'bounding_box': [point for point in paragraph.bounding_regions[0].polygon] if paragraph.bounding_regions else [],
                        'spans': paragraph.spans
                    }
                    paragraphs.append(paragraph_data)
            
            # OCR retorna apenas tokens puros - extração de campos é feita pela IA
            
            return {
                'fullText': full_text,
                'organizedLines': organized_lines,
                'paragraphs': paragraphs,
                'tables': [self._extract_table_data(table) for table in analyze_result.tables],
                'totalPages': len(analyze_result.pages),
                'totalLines': len(organized_lines),
                'totalWords': len(full_text.split()),
                'analysisType': 'prebuilt-read-v4',
                'documentType': self._detect_q88_type(full_text),
                'timestamp': time.strftime('%Y-%m-%dT%H:%M:%S')
            }
    
        except Exception as e:
            logger.error(f"❌ Erro no processamento OCR: {str(e)}")
            raise
    
    def _extract_table_data(self, table) -> Dict[str, Any]:
        """Extrai dados de uma tabela"""
        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append({
                    'content': cell.content,
                    'confidence': cell.confidence,
                    'row_index': cell.row_index,
                    'column_index': cell.column_index
                })
            rows.append(row_data)
        
        return {
            'row_count': table.row_count,
            'column_count': table.column_count,
            'cells': rows
        }
    
    def _is_important_q88_line(self, line_text: str) -> bool:
        """Verifica se uma linha contém informações importantes do Q88"""
        important_keywords = [
            'vessel', 'imo', 'flag', 'port', 'call', 'mmsi', 'builder', 'delivered', 'tonnage', 'gross',
            'net', 'dwt', 'loa', 'beam', 'classification', 'certificate'
        ]
        return any(keyword in line_text.lower() for keyword in important_keywords)
    
    def _detect_q88_type(self, text: str) -> str:
        """Detecta o tipo de documento Q88 baseado no conteúdo"""
        text_lower = text.lower()
        
        if 'intertanko' in text_lower:
                return 'Q88 (INTERTANKO)'
        elif 'q88' in text_lower:
            return 'Q88 (Generic)'
        else:
            return 'Unknown Q88'
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai texto de arquivo .docx usando python-docx
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx não está instalado! Instale com: pip install python-docx")
            return {
                'full_text': '',
                'organized_lines': [],
                'total_pages': 0,
                'total_lines': 0,
                'total_words': 0,
                'total_confidence_score': 0.0,
                'ocr_model_version': 'python-docx'
            }
        
        try:
            doc = Document(file_path)
            full_text = []
            organized_lines = []
            line_number = 0
            
            # Extrair texto de parágrafos
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
                    organized_lines.append({
                        'text': text,
                        'page_number': 1,  # .docx não tem conceito de páginas
                        'line_number': line_number,
                        'confidence': 1.0
                    })
                    line_number += 1
            
            # Extrair texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text = ' | '.join(row_text)
                        full_text.append(text)
                        organized_lines.append({
                            'text': text,
                            'page_number': 1,
                            'line_number': line_number,
                            'confidence': 1.0
                        })
                        line_number += 1
            
            full_text_str = '\n'.join(full_text)
            total_words = len(full_text_str.split())
            
            return {
                'full_text': full_text_str,
                'organized_lines': organized_lines,
                'total_pages': 1,
                'total_lines': line_number,
                'total_words': total_words,
                'total_confidence_score': 1.0,
                'ocr_model_version': 'python-docx'
            }
        except Exception as e:
            logger.error(f"Erro ao processar .docx: {str(e)}")
            return {
                'full_text': '',
                'organized_lines': [],
                'total_pages': 0,
                'total_lines': 0,
                'total_words': 0,
                'total_confidence_score': 0.0,
                'ocr_model_version': 'python-docx-error'
            }